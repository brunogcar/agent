"""
Write DOCX action handler.
"""

from __future__ import annotations

import re as _re
from pathlib import Path

from tools.file_ops.helpers import _safe_resolve
from tools.file_ops._registry import register_action

@register_action("file", "write_docx")
def _handle_write_docx(path: str = "", content: str = "", title: str = "") -> dict:
    """Write text to a DOCX file using python-docx."""
    p, err = _safe_resolve(path)
    if err:
        return {"status": "error", "error": err}
    if not content:
        return {"status": "error", "error": "content is required for write_docx"}
    if p.suffix.lower() != ".docx":
        p = p.with_suffix(".docx")

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        if title:
            t = doc.add_heading(title, level=0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for line in content.splitlines():
            stripped = line.strip()
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith(("- ", "* ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif _re.match(r"^\d+\.\s", stripped):
                doc.add_paragraph(_re.sub(r"^\d+\.\s", "", stripped),
                                  style="List Number")
            elif stripped in ("---", "***"):
                doc.add_paragraph("_" * 60)
            elif not stripped:
                doc.add_paragraph("")
            else:
                para = doc.add_paragraph()
                parts = _re.split(r"(\*\*[^*]+\*\*)", stripped)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        para.add_run(part[2:-2]).bold = True
                    else:
                        para.add_run(part)

        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        return {
            "status":     "success",
            "path":       str(p),
            "size":       p.stat().st_size,
            "paragraphs": len(doc.paragraphs),
        }
    except ImportError:
        return {"status": "error", "error": "python-docx not installed. Run: pip install python-docx"}
    except Exception as e:
        return {"status": "error", "error": f"DOCX write failed: {type(e).__name__}: {e}"}